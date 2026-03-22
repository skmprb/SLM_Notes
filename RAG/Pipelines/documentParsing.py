import re
from typing import Dict, List
from bs4 import BeautifulSoup
import PyPDF2
import pdfplumber
import fitz
import docx
import json
import pandas as pd

class DocumentParser:
    def parse_pdf_pypdf2(self, pdf_path: str) -> Dict:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            pages = []
            for i, page in enumerate(reader.pages):
                pages.append({
                    'page_num' : i+1,
                    'text': page.extract_text()
                })
        return {'source': pdf_path, 'pages': pages, 'type': 'pdf'}
    
    def parse_pdf_pdfplumber(self, pdf_path: str) -> Dict:
        with pdfplumber.open(pdf_path) as pdf:
            pages = []
            for i, page in enumerate(pdf.pages):
                pages.append({
                    'page_num' : i+1,
                    'text': page.extract_text(),
                    'tables': page.extract_tables()
                })
        return {'source': pdf_path, 'pages': pages, 'type': 'pdf'}
    
    def parse_pdf_pymupdf(self, pdf_path: str) -> Dict:
        doc = fitz.open(pdf_path)
        pages = []
        for i, page in enumerate(doc):
            pages.append({
                'page_num' : i+1,
                'text': page.get_text(),
                'images': len(page.get_images())
            })
        metadata = doc.metadata
        return {'source': pdf_path, 'pages': pages, 'type': 'pdf', 'metadata': metadata}
    
    def parse_html(self, html_content: str) -> Dict:
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for script in soup(["script", "style"]):
            script.decompose()
        
        title = soup.find('title').text if soup.find('title') else ''
        headings = [h.text.strip() for h in soup.find_all(['h1', 'h2', 'h3', 'h4'])]
        paragraphs = [p.text.strip() for p in soup.find_all('p')]
        links = [a.get('href') for a in soup.find_all('a', href = True)]
        return {
            'source': 'html_content',
            'title': title,
            'headings': headings,
            'paragraphs': paragraphs,
            'links': links,
            'full_text': soup.get_text(separator='\n', strip=True),
            'type': 'html'
        }
    
    def parse_docx(self, docx_path: str) -> Dict:
        doc = docx.Document(docx_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        tables = []
        for table in doc.tables:
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(table_data)
        
        return {
            'source': docx_path,
            'paragraphs': paragraphs,
            'tables': tables,
            'type': 'docx'
        }
        
    def parse_json(self, json_path: str) -> Dict:
        with open(json_path, 'r') as file:
            data = json.load(file)
        return {'source': json_path, 'data': data, 'type': 'json'}
    
    def parse_csv(self, csv_path: str) -> Dict:
        df = pd.read_csv(csv_path)
        return {
            'source': csv_path,
            'columns' : df.columns.tolist(),
            'rows': df.to_dict('records'),
            'type': 'csv',
            'text': df.to_string()
        }
    
    def parse_markdown(self, md_path: str) -> Dict:
        
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        headings = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
        
        code_blocks = re.findall(r'```(.*?)```', content)
        
        return{
            'source': md_path,
            'headings': headings,
            'code_blocks': code_blocks,
            'full_text': content,
            'type': 'markdown'
        }
    
    # --- Vision OCR Parsing Methods ---
    
    def parse_pdf_tesseract(self, pdf_path: str) -> Dict:
        """OCR for scanned PDFs using Tesseract + pdf2image"""
        from pdf2image import convert_from_path
        import pytesseract
        
        images = convert_from_path(pdf_path)
        pages = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(img)
            pages.append({'page_num': i + 1, 'text': text})
        return {'source': pdf_path, 'pages': pages, 'type': 'pdf', 'parser': 'tesseract_ocr'}
    
    def parse_image_tesseract(self, image_path: str) -> Dict:
        """OCR for images (PNG, JPG, TIFF) using Tesseract"""
        import pytesseract
        from PIL import Image
        
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return {'source': image_path, 'text': text, 'type': 'image', 'parser': 'tesseract_ocr'}
    
    def parse_with_textract(self, file_path: str, s3_bucket: str = None, s3_key: str = None) -> Dict:
        """AWS Textract - extracts text, tables, forms from documents/images"""
        import boto3
        
        client = boto3.client('textract')
        
        if s3_bucket and s3_key:
            response = client.detect_document_text(
                Document={'S3Object': {'Bucket': s3_bucket, 'Name': s3_key}}
            )
        else:
            with open(file_path, 'rb') as f:
                response = client.detect_document_text(Document={'Bytes': f.read()})
        
        lines = [block['Text'] for block in response['Blocks'] if block['BlockType'] == 'LINE']
        return {
            'source': s3_key or file_path,
            'text': '\n'.join(lines),
            'lines': lines,
            'type': 'textract',
            'parser': 'aws_textract'
        }
    
    def parse_with_gpt4_vision(self, image_path: str, prompt: str = "Extract all text from this image. Preserve structure and formatting.") -> Dict:
        """GPT-4 Vision - understands context, layout, handwriting"""
        from openai import OpenAI
        import base64
        
        client = OpenAI()
        with open(image_path, 'rb') as f:
            b64_image = base64.b64encode(f.read()).decode('utf-8')
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64_image}"}}
                ]
            }],
            max_tokens=4096
        )
        text = response.choices[0].message.content
        return {'source': image_path, 'text': text, 'type': 'image', 'parser': 'gpt4_vision'}
    
    def extract_text_from_parsed(self, parsed_data: Dict) -> str:
        
        doc_type = parsed_data.get('type')
        
        if doc_type == 'pdf':
            return '\n'.join([page['text'] for page in parsed_data['pages']])
        elif doc_type == 'html':
            return parsed_data['full_text']
        elif doc_type == 'docx':
            return '\n'.join(parsed_data['paragraphs'])
        elif doc_type == 'json':
            return json.dumps(parsed_data['data'], indent=2)
        elif doc_type == 'csv':
            return parsed_data['text']
        elif doc_type == 'markdown':
            return parsed_data['full_text']
        elif doc_type in ('image', 'textract'):
            return parsed_data.get('text', '')
        else:
            return str(parsed_data)
        
        